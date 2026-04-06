#!/usr/bin/env python3
"""
Generate professional PDF and DOCX from the Transparency Portal business case article.

Usage:
    python3 generate.py

Outputs:
    transparency-portal-business-case.pdf
    transparency-portal-business-case.docx
"""

import re
import os
from pathlib import Path

# --- PDF generation with reportlab ---
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor
from reportlab.lib.units import mm, cm
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable
)

# --- DOCX generation with python-docx ---
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


SCRIPT_DIR = Path(__file__).parent
MD_FILE = SCRIPT_DIR / "transparency-portal-business-case.md"
PDF_FILE = SCRIPT_DIR / "transparency-portal-business-case.pdf"
DOCX_FILE = SCRIPT_DIR / "transparency-portal-business-case.docx"

# Brand colors
MAGENTA = "#4F46E5"
DARK_BG = "#1C1C27"
COPPER = "#B87333"
SAND = "#F5F0E8"
DARK_TEXT = "#1A1A2E"
GRAY_TEXT = "#6B7280"

# ---------- Markdown Parser ----------

def parse_markdown(filepath):
    """Parse the markdown into structured blocks."""
    text = filepath.read_text(encoding="utf-8")
    blocks = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
            i += 1
            continue
        if line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
            i += 1
            continue
        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            blocks.append(("hr", ""))
            i += 1
            continue

        # Blockquote (pull quote)
        if line.startswith("> "):
            quote_text = line[2:].strip()
            i += 1
            while i < len(lines) and lines[i].startswith("> "):
                quote_text += " " + lines[i][2:].strip()
                i += 1
            # Remove surrounding *
            quote_text = quote_text.strip("*").strip()
            blocks.append(("quote", quote_text))
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and "---" in lines[i + 1]:
            headers = [c.strip() for c in line.split("|") if c.strip()]
            i += 2  # skip separator
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                row = [c.strip() for c in lines[i].split("|") if c.strip()]
                rows.append(row)
                i += 1
            blocks.append(("table", (headers, rows)))
            continue

        # Bold author line
        if line.startswith("**") and line.endswith("**"):
            blocks.append(("author", line.strip("*").strip()))
            i += 1
            continue

        # Regular paragraph
        para = line
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not lines[i].startswith(">") and not lines[i].startswith("|") and not lines[i].startswith("---") and not lines[i].startswith("**"):
            para += " " + lines[i]
            i += 1
        blocks.append(("para", para.strip()))

    return blocks


def clean_markdown_inline(text):
    """Convert inline markdown to reportlab XML or plain text."""
    # Bold
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    # Italic
    text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
    # Links: [text](url) -> text
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def clean_markdown_plain(text):
    """Strip all markdown formatting for DOCX."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'\[(.+?)\]\((.+?)\)', r'\1', text)
    return text


# ---------- PDF Generation ----------

def generate_pdf(blocks):
    doc = SimpleDocTemplate(
        str(PDF_FILE),
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    s_title = ParagraphStyle(
        "ArticleTitle", parent=styles["Title"],
        fontName="Helvetica-Bold", fontSize=22, leading=28,
        textColor=HexColor(DARK_TEXT), alignment=TA_LEFT,
        spaceAfter=6 * mm,
    )
    s_author = ParagraphStyle(
        "Author", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10, leading=14,
        textColor=HexColor(GRAY_TEXT), alignment=TA_LEFT,
        spaceAfter=4 * mm,
    )
    s_h2 = ParagraphStyle(
        "H2", parent=styles["Heading2"],
        fontName="Helvetica-Bold", fontSize=15, leading=20,
        textColor=HexColor(MAGENTA), spaceBefore=10 * mm, spaceAfter=4 * mm,
    )
    s_h3 = ParagraphStyle(
        "H3", parent=styles["Heading3"],
        fontName="Helvetica-Bold", fontSize=12, leading=16,
        textColor=HexColor(DARK_TEXT), spaceBefore=6 * mm, spaceAfter=3 * mm,
    )
    s_body = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10, leading=15,
        textColor=HexColor(DARK_TEXT), alignment=TA_JUSTIFY,
        spaceAfter=3 * mm,
    )
    s_quote = ParagraphStyle(
        "PullQuote", parent=styles["Normal"],
        fontName="Helvetica-Oblique", fontSize=11, leading=16,
        textColor=HexColor(COPPER), alignment=TA_LEFT,
        leftIndent=15 * mm, rightIndent=15 * mm,
        spaceBefore=6 * mm, spaceAfter=6 * mm,
        borderColor=HexColor(COPPER), borderWidth=0,
        borderPadding=0,
    )
    s_source = ParagraphStyle(
        "Source", parent=styles["Normal"],
        fontName="Helvetica", fontSize=8, leading=11,
        textColor=HexColor(GRAY_TEXT), spaceAfter=1.5 * mm,
    )

    story = []

    for btype, content in blocks:
        if btype == "h1":
            story.append(Paragraph(clean_markdown_inline(content), s_title))
        elif btype == "author":
            story.append(Paragraph(content, s_author))
        elif btype == "h2":
            story.append(Paragraph(clean_markdown_inline(content), s_h2))
        elif btype == "h3":
            story.append(Paragraph(clean_markdown_inline(content), s_h3))
        elif btype == "hr":
            story.append(Spacer(1, 3 * mm))
            story.append(HRFlowable(
                width="100%", thickness=0.5,
                color=HexColor("#E5E7EB"), spaceAfter=3 * mm
            ))
        elif btype == "quote":
            story.append(Paragraph(
                f'<i>"{clean_markdown_inline(content)}"</i>', s_quote
            ))
        elif btype == "table":
            headers, rows = content
            table_data = [headers] + rows
            # Clean markdown from cells
            table_data = [
                [clean_markdown_inline(cell) for cell in row]
                for row in table_data
            ]
            # Convert to Paragraphs for wrapping
            col_count = len(headers)
            avail_width = A4[0] - 5 * cm
            col_width = avail_width / col_count

            cell_style = ParagraphStyle(
                "Cell", fontName="Helvetica", fontSize=8.5,
                leading=11, textColor=HexColor(DARK_TEXT),
            )
            header_style = ParagraphStyle(
                "CellHeader", fontName="Helvetica-Bold", fontSize=8.5,
                leading=11, textColor=HexColor("#FFFFFF"),
            )

            formatted = []
            for ri, row in enumerate(table_data):
                formatted_row = []
                for cell in row:
                    st = header_style if ri == 0 else cell_style
                    formatted_row.append(Paragraph(cell, st))
                formatted.append(formatted_row)

            t = Table(formatted, colWidths=[col_width] * col_count)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), HexColor(MAGENTA)),
                ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#FFFFFF")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#E5E7EB")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#FFFFFF"), HexColor("#F9FAFB")]),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(Spacer(1, 2 * mm))
            story.append(t)
            story.append(Spacer(1, 3 * mm))
        elif btype == "para":
            cleaned = clean_markdown_inline(content)
            # Check if it's a sources section item
            if cleaned.startswith("- "):
                story.append(Paragraph(cleaned[2:], s_source))
            else:
                story.append(Paragraph(cleaned, s_body))

    doc.build(story)
    print(f"PDF generated: {PDF_FILE}")


# ---------- DOCX Generation ----------

def generate_docx(blocks):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    for btype, content in blocks:
        text = clean_markdown_plain(content) if isinstance(content, str) else content

        if btype == "h1":
            p = doc.add_heading(text, level=1)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                run.font.size = Pt(22)

        elif btype == "author":
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
                run.italic = True

        elif btype == "h2":
            p = doc.add_heading(text, level=2)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xE2, 0x00, 0x74)
                run.font.size = Pt(15)

        elif btype == "h3":
            p = doc.add_heading(text, level=3)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                run.font.size = Pt(12)

        elif btype == "hr":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a thin border line
            run = p.add_run("_" * 80)
            run.font.color.rgb = RGBColor(0xE5, 0xE7, 0xEB)
            run.font.size = Pt(2)

        elif btype == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.right_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(f'"{clean_markdown_plain(text)}"')
            run.italic = True
            run.font.color.rgb = RGBColor(0xB8, 0x73, 0x33)  # Copper
            run.font.size = Pt(11)

        elif btype == "table":
            headers, rows = content
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.style = "Table Grid"

            # Header row
            for ci, header in enumerate(headers):
                cell = table.rows[0].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(clean_markdown_plain(header))
                run.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                from docx.oxml.ns import qn
                shading = cell._element.get_or_add_tcPr()
                shading_elem = shading.makeelement(qn("w:shd"), {
                    qn("w:fill"): "E20074",
                    qn("w:val"): "clear",
                })
                shading.append(shading_elem)

            # Data rows
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    cell = table.rows[ri + 1].cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(clean_markdown_plain(cell_text))
                    run.font.size = Pt(8.5)

            doc.add_paragraph()  # spacer

        elif btype == "para":
            plain = clean_markdown_plain(content) if isinstance(content, str) else str(content)
            if plain.startswith("- "):
                # Source list item
                p = doc.add_paragraph(plain[2:])
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            else:
                p = doc.add_paragraph(plain)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(str(DOCX_FILE))
    print(f"DOCX generated: {DOCX_FILE}")


# ---------- Main ----------

if __name__ == "__main__":
    print(f"Reading: {MD_FILE}")
    blocks = parse_markdown(MD_FILE)
    print(f"Parsed {len(blocks)} blocks")
    generate_pdf(blocks)
    generate_docx(blocks)
    print("Done.")
